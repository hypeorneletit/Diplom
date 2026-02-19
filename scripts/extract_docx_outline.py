from __future__ import annotations

import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document


def _force_utf8_stdout() -> None:
    # Make console output stable on Windows terminals.
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass


def is_heading_style(style_name: str) -> bool:
    s = style_name.lower()
    return s.startswith("heading") or s.startswith("title") or "заголов" in s


def looks_like_numbered_heading(text: str) -> bool:
    # Examples: "1 ...", "1.2 ...", "2.3.4 ...", "Глава 1 ...", "Раздел 2 ..."
    t = text.strip()
    if re.match(r"^(глава|раздел)\s+\d+\b", t, flags=re.IGNORECASE):
        return True
    if re.match(r"^\d+(?:[\.,]\d+){0,4}[\.)]\s+\S+", t):
        return True
    return False


def main() -> int:
    _force_utf8_stdout()

    if len(sys.argv) < 2:
        print("Usage: extract_docx_outline.py <path-to-docx | path-to-txt-with-docx-path>")
        return 2

    arg = Path(sys.argv[1]).expanduser()
    if arg.suffix.lower() == ".txt" and arg.exists():
        docx_path_str = arg.read_text(encoding="utf-8", errors="strict").strip().strip("\ufeff")
        docx_path = Path(docx_path_str).expanduser().resolve()
    else:
        docx_path = arg.resolve()

    doc = Document(docx_path)

    out_dir = Path(__file__).parent
    outline_txt = out_dir / "diplom_outline.txt"
    full_txt = out_dir / "diplom_fulltext.txt"

    para_count = len(doc.paragraphs)

    style_counter = Counter()
    headings: list[tuple[str, str]] = []
    numbered_candidates: list[str] = []

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style = (p.style.name if p.style else "").strip()
        style_counter[style] += 1

        if style and is_heading_style(style):
            headings.append((style, text))
        elif looks_like_numbered_heading(text):
            numbered_candidates.append(text)

    print(f"PARA_COUNT\t{para_count}")
    print(f"NONEMPTY_PARA\t{sum(style_counter.values())}")
    print(f"HEADING_STYLE_PARA\t{len(headings)}")

    print("\nTOP_STYLES:")
    for style, cnt in style_counter.most_common(12):
        label = style if style else "(no-style)"
        print(f"- {label}: {cnt}")

    outline_lines: list[str] = []
    if headings:
        print("\nOUTLINE_FROM_STYLES:")
        for style, text in headings[:120]:
            line = f"- [{style}] {text}"
            outline_lines.append(line)
            print(line)
        if len(headings) > 120:
            outline_lines.append("... truncated ...")
            print("... truncated ...")
    else:
        print("\nOUTLINE_FROM_NUMBERING (no heading styles detected):")
        for text in numbered_candidates[:120]:
            line = f"- {text}"
            outline_lines.append(line)
            print(line)
        if len(numbered_candidates) > 120:
            outline_lines.append("... truncated ...")
            print("... truncated ...")

    # Persist extracted text so it can be analyzed without terminal encoding issues.
    outline_txt.write_text("\n".join(outline_lines) + "\n", encoding="utf-8")
    full_txt.write_text(
        "\n".join((p.text or "").rstrip() for p in doc.paragraphs) + "\n",
        encoding="utf-8",
    )
    print(f"\nWROTE_OUTLINE\t{outline_txt}")
    print(f"WROTE_FULLTEXT\t{full_txt}")

    # Keyword scan to help tailor ideas.
    keywords = [
        "vr",
        "unity",
        "oculus",
        "meta",
        "quest",
        "steamvr",
        "openxr",
        "комнат",
        "сценар",
        "метрик",
        "энерг",
        "диспетчер",
        "оператор",
        "монитор",
        "интерактив",
        "телепорт",
        "жест",
        "голос",
        "alarm",
        "тревог",
        "kpi",
        "ui",
        "ux",
        "эргоном",
        "время реакции",
        "ошиб",
    ]
    hits = Counter()
    for p in doc.paragraphs:
        t = (p.text or "").lower()
        for k in keywords:
            if k in t:
                hits[k] += 1

    nonzero = [(k, v) for k, v in hits.items() if v]
    nonzero.sort(key=lambda kv: (-kv[1], kv[0]))

    print("\nKEYWORD_HITS:")
    if not nonzero:
        print("(none)")
    else:
        for k, v in nonzero[:40]:
            print(f"- {k}: {v}")
        if len(nonzero) > 40:
            print("... truncated ...")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
